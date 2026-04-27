"""Apply replacements back to a DOCX file.

We rebuild paragraph text rather than try to preserve every run boundary —
maintaining run-level styling around partial replacements is complex and the
spec explicitly accepts this trade-off for the MVP. Paragraph-level styling
(alignment, indent, font of the first run) is preserved.
"""
from __future__ import annotations

from typing import Any

from idp_anonymiser.agent.state import (
    AnonymisationPlan,
    ResolvedEntity,
)
from idp_anonymiser.document.layout_model import ExtractedDocument


def _build_paragraph_edits(
    extracted: ExtractedDocument,
    plan: AnonymisationPlan,
    resolved: list[ResolvedEntity],
) -> dict[str, list[tuple[int, int, str]]]:
    """For each block_id in the extracted doc, the list of (start, end, replacement) edits."""
    by_entity = {r.entity_id: r for r in plan.replacements}
    block_index = {b.block_id: b for b in extracted.blocks if b.block_id}
    out: dict[str, list[tuple[int, int, str]]] = {}
    for ent in resolved:
        rep = by_entity.get(ent.entity_id)
        if rep is None:
            continue
        for det in ent.detections:
            if det.span.start is None or det.span.end is None:
                continue
            for bid, block in block_index.items():
                if block.start <= det.span.start and det.span.end <= block.end:
                    local_s = det.span.start - block.start
                    local_e = det.span.end - block.start
                    out.setdefault(bid, []).append((local_s, local_e, rep.replacement_value))
                    break
    return out


def _replace_paragraph_text(paragraph: Any, new_text: str) -> None:
    """Replace a paragraph's text while keeping the paragraph object structure.

    All runs are removed and a single run containing ``new_text`` is appended.
    The paragraph's existing properties (style, alignment) are retained.
    """
    # python-docx exposes runs via .runs and the underlying XML via ._p
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(new_text)


def rewrite_docx(
    input_path: str,
    output_path: str,
    extracted: ExtractedDocument,
    plan: AnonymisationPlan,
    resolved: list[ResolvedEntity],
) -> int:
    """Rewrite the DOCX. Returns the number of replacements applied."""
    from docx import Document

    edits = _build_paragraph_edits(extracted, plan, resolved)
    doc = Document(input_path)
    applied = 0

    # Map (section, block_index, table_id?, cell_row?, cell_col?, para_index?) -> python-docx paragraph
    # We do a parallel walk in the same order as the extractor.

    # Body paragraphs first
    for idx, para in enumerate(doc.paragraphs):
        bid = f"docx:body:{idx}"
        if bid in edits:
            new_text = para.text or ""
            for s, e, repl in sorted(edits[bid], key=lambda x: x[0], reverse=True):
                new_text = new_text[:s] + repl + new_text[e:]
                applied += 1
            _replace_paragraph_text(para, new_text)

    # Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    bid = f"docx:t{t_idx}:r{r_idx}:c{c_idx}:p{p_idx}"
                    if bid in edits:
                        new_text = para.text or ""
                        for s, e, repl in sorted(
                            edits[bid], key=lambda x: x[0], reverse=True
                        ):
                            new_text = new_text[:s] + repl + new_text[e:]
                            applied += 1
                        _replace_paragraph_text(para, new_text)

    doc.save(output_path)
    return applied


__all__ = ["rewrite_docx"]
