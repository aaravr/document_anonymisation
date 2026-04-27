"""Extract paragraphs and tables from a DOCX file via python-docx."""
from __future__ import annotations

from idp_anonymiser.document.layout_model import (
    DocxParagraph,
    ExtractedDocument,
    ExtractedTextBlock,
)
from idp_anonymiser.document.loader import LoadedDocument


def extract_docx(loaded: LoadedDocument) -> ExtractedDocument:
    from docx import Document

    doc = Document(str(loaded.path))
    parts: list[str] = []
    blocks: list[ExtractedTextBlock] = []
    paragraphs: list[DocxParagraph] = []
    cursor = 0

    # Body paragraphs
    for idx, para in enumerate(doc.paragraphs):
        text = para.text or ""
        paragraphs.append(
            DocxParagraph(section="body", block_index=idx, text=text)
        )
        if not text.strip():
            cursor += 1  # placeholder for blank line
            parts.append("\n")
            continue
        start = cursor
        parts.append(text)
        cursor += len(text)
        end = cursor
        parts.append("\n")
        cursor += 1
        blocks.append(
            ExtractedTextBlock(
                text=text,
                start=start,
                end=end,
                block_id=f"docx:body:{idx}",
                metadata={"section": "body", "block_index": idx},
            )
        )

    # Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text or ""
                    paragraphs.append(
                        DocxParagraph(
                            section="table",
                            block_index=p_idx,
                            text=text,
                            table_id=f"t{t_idx}",
                            cell_row=r_idx,
                            cell_col=c_idx,
                        )
                    )
                    if not text.strip():
                        continue
                    start = cursor
                    parts.append(text)
                    cursor += len(text)
                    end = cursor
                    parts.append("\n")
                    cursor += 1
                    blocks.append(
                        ExtractedTextBlock(
                            text=text,
                            start=start,
                            end=end,
                            block_id=f"docx:t{t_idx}:r{r_idx}:c{c_idx}:p{p_idx}",
                            metadata={
                                "section": "table",
                                "table_index": t_idx,
                                "row": r_idx,
                                "column": c_idx,
                                "para_index": p_idx,
                            },
                        )
                    )
    return ExtractedDocument(
        flat_text="".join(parts),
        blocks=blocks,
        docx_paragraphs=paragraphs,
        metadata={"format": "docx"},
    )
