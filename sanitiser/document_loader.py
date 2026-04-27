"""Document loader: TXT, DOCX, PDF (searchable). Returns page-level text."""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional


@dataclass
class PageText:
    page_index: int      # 0-indexed
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class LoadedDoc:
    doc_id: str
    input_path: Path
    format: str  # "txt" | "docx" | "pdf"
    pages: list[PageText]
    raw_handle: Any = None  # native handle (e.g. fitz.Document) for downstream rewriting


def detect_format(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt", ".log", ".md"}:
        return "txt"
    if ext == ".docx":
        return "docx"
    if ext == ".pdf":
        return "pdf"
    raise ValueError(f"Unsupported file extension: {ext} (expected .txt, .docx, or .pdf)")


def _load_txt(path: Path, doc_id: str) -> LoadedDoc:
    text = path.read_text(encoding="utf-8", errors="replace")
    # Treat the whole TXT as a single "page".
    return LoadedDoc(doc_id=doc_id, input_path=path, format="txt", pages=[PageText(0, text)])


def _load_docx(path: Path, doc_id: str) -> LoadedDoc:
    from docx import Document
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    text = "\n".join(parts)
    return LoadedDoc(
        doc_id=doc_id, input_path=path, format="docx",
        pages=[PageText(0, text)], raw_handle=doc,
    )


def _load_pdf(path: Path, doc_id: str) -> LoadedDoc:
    import fitz  # PyMuPDF
    pdf = fitz.open(str(path))
    pages: list[PageText] = []
    for i in range(pdf.page_count):
        page = pdf.load_page(i)
        text = page.get_text("text") or ""
        pages.append(PageText(i, text, {"page_count": pdf.page_count}))
    return LoadedDoc(
        doc_id=doc_id, input_path=path, format="pdf",
        pages=pages, raw_handle=pdf,
    )


def load(path: str | Path, doc_id: Optional[str] = None) -> LoadedDoc:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(p)
    fmt = detect_format(p)
    did = doc_id or p.stem
    if fmt == "txt":
        return _load_txt(p, did)
    if fmt == "docx":
        return _load_docx(p, did)
    if fmt == "pdf":
        return _load_pdf(p, did)
    raise AssertionError(f"unreachable: {fmt}")
