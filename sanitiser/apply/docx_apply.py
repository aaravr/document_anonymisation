"""Apply replacements to a DOCX file paragraph-by-paragraph.

Like the IDP rewriter, paragraph-level styling is preserved but per-run
formatting around partial replacements is rebuilt as a single run.
"""
from __future__ import annotations

import re
from sanitiser.state import Detection, ReplacementRecord, CanonicalEntity
from sanitiser.apply.text_apply import _pick_variant
from sanitiser.resolve.normaliser import normalise


def rewrite_docx(input_path: str, output_path: str,
                  detections_by_page: dict[int, list[Detection]],
                  entity_lookup, *, document_id: str) -> tuple[int, list[ReplacementRecord]]:
    from docx import Document
    doc = Document(input_path)

    audit: list[ReplacementRecord] = []
    applied = 0

    # We re-derive the flat paragraph text and apply per-paragraph string substitutions.
    # The detections we received reference the page-level flat text we extracted
    # earlier; here we just use the per-mention surface form to do a literal
    # search-and-replace per paragraph.

    # Build a global ordered replacement list: deterministic by detection order.
    flat_replacements: list[tuple[str, str, Detection, CanonicalEntity]] = []
    for page_idx, dets in detections_by_page.items():
        for d in dets:
            ent = entity_lookup(d.entity_type, normalise(d.text, d.entity_type))
            if ent is None:
                continue
            repl = _pick_variant(d.text, ent)
            if d.text.isupper() and any(c.isalpha() for c in d.text):
                repl = repl.upper()
            flat_replacements.append((d.text, repl, d, ent))

    # Sort longer originals first to avoid replacing an abbreviation inside a full name.
    flat_replacements.sort(key=lambda x: -len(x[0]))

    def _apply_to_paragraph(para):
        nonlocal applied
        text = para.text or ""
        new_text = text
        for original, replacement, det, ent in flat_replacements:
            if original in new_text:
                new_text, count = _replace_word_boundary(new_text, original, replacement)
                if count > 0:
                    applied += count
                    audit.append(ReplacementRecord(
                        document_id=document_id, page=0, chunk=None,
                        entity_type=det.entity_type, original=original, replacement=replacement,
                        canonical_id=ent.canonical_id, detectors=[det.detector],
                        confidence=det.confidence, start=0, end=len(original),
                        reason=det.detector + " match in DOCX paragraph",
                    ))
        if new_text != text:
            for run in list(para.runs):
                run._element.getparent().remove(run._element)
            para.add_run(new_text)

    for p in doc.paragraphs:
        _apply_to_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _apply_to_paragraph(p)
    doc.save(output_path)
    return applied, audit


def _replace_word_boundary(haystack: str, needle: str, replacement: str) -> tuple[str, int]:
    """Replace ``needle`` in ``haystack`` honoring word boundaries when the
    needle starts/ends with an alphanumeric. Returns (new_text, count)."""
    if not needle:
        return haystack, 0
    pat = re.compile(r"\b" + re.escape(needle) + r"\b")
    new, n = pat.subn(replacement, haystack)
    return new, n
